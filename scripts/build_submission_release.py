#!/usr/bin/env python3
"""Build the local-only LightGuard Submission Release v1.0 report package."""

from __future__ import annotations

import hashlib
import html
import subprocess
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "release-v1.0"
FIG = OUT / "figures"

INK, PAPER, CANVAS = "#102A43", "#FFFDF8", "#F4F1EA"
SIGNAL, CAUTION, RED = "#0F766E", "#D97706", "#B42318"
LINE, MUTED = "#D8D4CA", "#52606D"
FONT = Path("/Library/Fonts/NanumBarunGothic.ttf")
FONT_BOLD = Path("/Library/Fonts/NanumBarunGothicBold.ttf")
FONT_DISPLAY = Path("/Library/Fonts/NanumSquareExtraBold.ttf")

DEFINITION = (
    "LightGuard는 가로등 고장을 자동 확정하는 시스템이 아니라, 기존 AMI의 실제 전력 흐름을 "
    "독립적으로 확인해 이상징후를 찾고, 반복 유지관리 이력과 결합해 담당자의 원격확인·현장점검 "
    "우선순위를 지원하는 서비스입니다."
)

METRICS = [
    ("204개", "수영구 분전함", "자산 master 기준"),
    ("4,239등", "연결 가로등", "수영구 자산 집계"),
    ("488.44 kW", "추정 정격부하", "자산 정규화 결과"),
    ("46개", "통제 시나리오 탐지", "실제 고장 정확도 아님"),
    ("101,843건", "대구 고장등 관리", "운영 workload 근거"),
    ("920 / 981", "울산 안전 위치연결", "13 모호·48 미연결"),
]

SECTIONS = [
    (
        "현장의 문제는 고장 확정이 아니라 확인 순서입니다",
        "가로등 운영자는 민원·일상점검·직원신고와 원격제어 상태를 함께 보지만, 실제 전력 흐름을 독립적으로 확인할 근거가 부족합니다.",
        [
            "가로등 이상은 시민 안전과 야간 이동에 연결되지만 모든 분전함과 등기구를 같은 빈도로 현장 점검하기는 어렵습니다. 기존 원격제어는 명령과 통신 상태를 보여주는 데 강점이 있으나, 명령 이후 실제 전력 사용이 기대 범위에 들어왔는지는 별도 계측 근거가 필요합니다.",
            "대구의 공개 고장등 관리자료 101,843건은 일상점검·직원신고·민원신고가 함께 작동하는 대규모 확인 업무를 보여줍니다. LightGuard는 이 업무를 대체하지 않고 먼저 확인할 대상을 좁힙니다.",
        ],
        "운영 병목은 자동 고장확정 모델의 부재보다 확인 대상의 우선순위가 불명확하다는 데 있습니다.",
    ),
    (
        "AMI를 가로등 제어의 두 번째 확인자로 사용합니다",
        "제어 명령과 독립적인 전력 계측으로 예상 점등시간·정격부하와 실제 사용량의 차이를 찾습니다.",
        [
            "LightGuard 객체는 분전함 → 자산정보 → 예상 점등시간 → 예상 정격부하 → AMI 실측 → 이상근거 → 점검우선순위 순서로 고정했습니다.",
            "부분소등처럼 예상보다 낮은 부하, 일출 이후에도 소비가 지속되는 신호를 찾은 뒤 데이터 품질 검토, 원격관찰, 현장점검 후보로 분기합니다. 최종 고장 여부와 원인은 담당자가 확정합니다.",
        ],
        "LightGuard는 고장 판정기가 아니라 제어 상태를 실제 전력 흐름으로 교차 확인하는 AMI second checker입니다.",
    ),
    (
        "서비스는 담당자의 다음 행동을 중심으로 설계했습니다",
        "이상징후 목록 → 이유 확인 → 운영이력 → 원격관찰 또는 현장점검 흐름입니다.",
        [
            "첫 화면은 연구 버전과 내부 등급이 아니라 오늘 확인할 자산을 보여줍니다. 담당자는 예상 부하, 실측 신호, 지속시간과 판정근거를 보고 과거 반복 이력과 backlog 맥락을 확인합니다.",
            "데이터 결측은 DATA_QUALITY_REVIEW, 약한 신호는 REMOTE_MONITOR, 재현 가능하고 설명 가능한 신호는 FIELD_INSPECTION_CANDIDATE로 제시해 과잉 출동을 막습니다.",
        ],
        "모델 출력의 종착점은 고장 라벨이 아니라 담당자가 수행할 수 있는 다음 확인 행동입니다.",
    ),
    (
        "수영구 204개 분전함을 앱이 쓰는 객체로 고정했습니다",
        "분전함 204개, 가로등 4,239등, 추정 정격부하 488.44 kW를 공간·천문·AMI 맥락과 결합합니다.",
        [
            "수영구 자산자료를 분전함 단위로 정규화하고 2026년 일출·일몰 및 시민박명, 대응 기상관측·예보 지점, 공간분포와 연결했습니다.",
            "실제 수영구 AMI 부재는 숨기지 않았습니다. 정격 3.4 kW 자산의 20% 부분소등, 일출 후 90분 점등 지속 같은 신호를 실제 자산에 주입해 판정 경로를 재현합니다. 이는 규칙 동작 검증이며 실제 현장 정확도가 아닙니다.",
        ],
        "실제 AMI 부재 구간은 시나리오 검증으로 분리하고 실제 성능 주장으로 전환하지 않습니다.",
    ),
    (
        "알고리즘은 SIGNAL과 OPERATIONS 두 층으로 분리됩니다",
        "전력 신호 탐지와 유지관리 우선순위를 분리해 지역별 가용 데이터 차이를 흡수합니다.",
        [
            "SIGNAL layer는 시민박명을 포함한 기대 점등창, 정격부하, AMI 관측치의 관계를 시간·품질·지속시간과 함께 평가합니다.",
            "OPERATIONS layer는 접수 당시 관측 가능했던 접수경로·구청·계절·과거 30/90/365일 사건·이전 장기처리·열린 사건·start-of-day backlog를 사용합니다. 같은 날 순서를 만들거나 미래 처리일을 가져오는 leakage는 금지합니다.",
        ],
        "전국 동일 모델을 주장하지 않고 공통 데이터 계약과 지역별 layer 조합으로 확장합니다.",
    ),
    (
        "검증은 탐지 가능성·외부 재현·운영 유용성을 분리했습니다",
        "서로 다른 근거가 답하는 질문을 섞지 않아 성능 수치가 과장된 운영효과로 번지지 않게 했습니다.",
        [
            "통제 시나리오는 알려진 이상을 Detector가 재현하는지, 실제 공모전 AMI 사례는 신호가 현실 계측에서 관찰 가능한지 확인합니다. 외부 benchmark 실패도 조건부 한계로 보존했습니다.",
            "대구·부여·울산 자료는 AMI 정확도의 대체 검증이 아닙니다. 실제 유지관리 workload, 반복 사건, 처리 흐름과 보수적 위치 연결을 보여주는 OPERATIONS 근거입니다.",
        ],
        "실험 성능, 현실 신호 가능성, 운영 활용 근거를 하나의 정확도 숫자로 합치지 않습니다.",
    ),
    (
        "결과는 제출 가능한 근거 사슬을 이룹니다",
        "SR-A, 독립 QA PASS, 32개 release artifact 해시 일치와 metric·claim·rubric registry 및 금지주장 검사를 통과했습니다.",
        [
            "v0.21 release는 수치와 주장의 출처를 registry로 연결하고 독립 QA와 artifact 해시를 고정했습니다. 이는 모든 지역에서 동일 성능을 낸다는 뜻이 아니라 문서와 앱이 검증 경계 안에서 같은 이야기를 한다는 뜻입니다.",
            "수영구 자산 객체, 46개 통제 시나리오, 공모전 AMI 사례, 대구 101,843건, 부여 3,437건, 울산 1,060건이 SIGNAL–OPERATIONS–PRODUCT 사슬을 구성합니다.",
        ],
        "현재의 강점은 단일 최고 정확도가 아니라 주장과 근거가 추적 가능한 제출용 제품 체계입니다.",
    ),
    (
        "실제 운영자료는 확인 대상을 좁히는 가치를 보여줍니다",
        "대구·부여·울산의 독립 자료는 지자체마다 다른 유지관리 업무를 공통 우선순위 문제로 연결합니다.",
        [
            "대구 안전점검 105,449건과 수선·자재 145,365건은 유지관리 부담을 보강하지만 개별 고장과 확정 join key가 없으면 비용절감 계산에 사용하지 않습니다.",
            "울산 위치자료는 981개 자산 중 920개를 안전하게 연결했고 13개는 모호, 48개는 미연결로 남겼습니다. 불확실한 join을 버리지 않고 상태로 관리합니다.",
        ],
        "유형효과는 민원 감소율이 아니라 대규모 workload에서 원격확인·현장점검 후보를 좁히는 운영지원 효과입니다.",
    ),
    (
        "주 사용자는 가로등 운영·유지관리 담당자입니다",
        "담당자의 판단을 빠르게 만드는 읽기 전용 보조 도구로 시작합니다.",
        [
            "지자체 도로조명 담당자와 시설관리기관 관제·유지보수 담당자가 1차 사용자입니다. 관리자는 workload를 보고 일정을 배분하고 현장 인력은 자산 위치·예상 부하·이상근거를 확인합니다.",
            "도입 초기에는 주간 점검목록과 shadow mode로 운영합니다. 실제 확인 결과를 feedback label로 축적해 지역별 threshold와 분기 규칙을 교정합니다.",
        ],
        "자동화의 목표는 담당자를 제거하는 것이 아니라 제한된 점검 시간을 설명 가능한 후보에 집중시키는 것입니다.",
    ),
    (
        "기존 원격제어를 교체하지 않고 계측 근거를 더합니다",
        "센서 추가 설치보다 기존 AMI와 자산·천문·운영자료를 우선 재사용합니다.",
        [
            "기존 원격제어는 스위칭·통신·상태관리의 중심으로 유지됩니다. LightGuard는 제어권 없이 실제 전력 흐름이 자산과 시간 맥락에 맞는지 확인하는 보조층입니다.",
            "등기구별 센서는 세밀하지만 설치·통신·유지보수 비용이 발생합니다. 분전함 AMI가 있는 환경에 먼저 적용하고 판별이 어려운 구간에만 추가 센서 도입 근거를 제공합니다.",
        ],
        "차별점은 제어 대체가 아니라 기존 인프라 사이의 검증 공백을 낮은 추가설비 부담으로 메우는 것입니다.",
    ),
    (
        "경제효과는 실제 pilot에서 측정할 지표로 남깁니다",
        "현재 자료만으로 비용절감·민원감소 금액을 확정하지 않습니다.",
        [
            "현재 공개자료에는 인력 수, 출동비, 실제 고장원인, 사용 전후 무작위 비교가 없습니다. 따라서 연간 절감액과 민원 감소율을 주장하지 않습니다.",
            "pilot에서는 후보 대비 원격 해소율, 현장점검 적중률, 사건당 확인시간, 반복 방문률, 장기 backlog, 데이터 품질 보류율을 사전 정의합니다.",
        ],
        "경제성은 홍보성 추정치가 아니라 운영 로그로 사후 검증할 pilot KPI입니다.",
    ),
    (
        "다음 단계는 실제 AMI를 연결한 제한적 현장 실증입니다",
        "새 모델링보다 Gold/Silver label과 현장 feedback을 먼저 확보합니다.",
        [
            "1단계는 실제 분전함 AMI와 자산 master를 읽기 전용으로 연결해 4~8주 shadow monitoring을 수행하는 것입니다. 2단계는 담당자 확인 결과를 구조화해 지역별 threshold를 보정합니다.",
            "3단계는 계절과 지역을 넓혀 SIGNAL 안정성과 OPERATIONS layer의 local calibration을 분리 평가합니다. 원시 계측과 개인정보는 최소수집하고 release artifact는 해시와 registry로 관리합니다.",
        ],
        "확장은 전국 동일 모델이 아니라 공통 계약, 지역별 보정, 사람의 확인 결과가 순환하는 방식으로 진행합니다.",
    ),
]

DISCLAIMERS = [
    "고장 여부와 원인을 자동 확정하지 않습니다.",
    "46개 시나리오 탐지는 실제 현장 정확도·재현율이 아닙니다.",
    "운영자료는 AMI 성능의 대체 검증이 아닙니다.",
    "민원 감소율과 비용절감액을 현재 결과로 확정하지 않습니다.",
    "한 모델이 전국에 무보정 적용된다고 주장하지 않습니다.",
]

SOURCES = [
    "공식 공모전 접수서류 및 평가자료: official_docs/ (로컬 비추적)",
    "LightGuard v0.21 metric·claim·rubric registry 및 독립 QA 산출물",
    "수영구 가로등 분전함 자산, 2026 천문·기상 context, 공모전 AMI 자료",
    "대구공공시설관리공단 고장등·위치·안전점검·수선자재·공사 자료",
    "부여군 고장·운영자료, 울산광역시 남구 접수·처리 및 위치자료",
]


def pil_font(size: int, bold: bool = False, display: bool = False):
    return ImageFont.truetype(str(FONT_DISPLAY if display else FONT_BOLD if bold else FONT), size)


def build_figures() -> list[Path]:
    FIG.mkdir(parents=True, exist_ok=True)
    paths = []
    specs = [
        (
            "01_service_flow.png",
            "판정이 아니라 다음 확인 행동으로",
            [("예상", "자산·정격부하\n일몰·시민박명"), ("관측", "AMI 전력 흐름\n품질·지속시간"), ("설명", "이상근거\n반복 운영이력"), ("행동", "원격관찰\n현장점검 후보")],
        ),
        (
            "02_evidence_layers.png",
            "서로 다른 근거는 서로 다른 질문에 답합니다",
            [("SIGNAL", "예상 대비 AMI"), ("PLAUSIBILITY", "현실 신호 가능성"), ("OPERATIONS", "반복·backlog"), ("PRODUCT", "담당자의 다음 행동")],
        ),
        (
            "03_claim_boundaries.png",
            "제출 가능한 주장은 검증 경계 안에 있습니다",
            [("확인", "탐지 규칙·현실 신호·운영 workload"), ("보류", "자동 고장확정·전국 정확도·비용절감")],
        ),
        (
            "04_validation_performance.png",
            "제출 전 검증 gate는 모두 통과했습니다",
            [("46 / 46", "통제 시나리오 탐지"), ("32 / 32", "release 해시 일치"), ("PASS", "독립 QA"), ("PASS", "금지주장 검사")],
        ),
        (
            "05_usage_example.png",
            "담당자는 네 단계로 점검 대상을 확인합니다",
            [("목록", "우선순위 #1"), ("이유", "20% 부분소등 신호"), ("이력", "반복 접수·backlog"), ("행동", "원격관찰 또는 현장점검")],
        ),
    ]
    for filename, title, cards in specs:
        image = Image.new("RGB", (1600, 900), CANVAS)
        draw = ImageDraw.Draw(image)
        draw.text((85, 65), title, font=pil_font(52, display=True), fill=INK)
        draw.text((87, 132), "LightGuard Submission Release v1.0 · 수치와 한계는 본문 참조", font=pil_font(22), fill=MUTED)
        width = (1430 - 30 * (len(cards) - 1)) // len(cards)
        for i, (head, body) in enumerate(cards):
            x = 85 + i * (width + 30)
            draw.rounded_rectangle((x, 260, x + width, 700), radius=28, fill=PAPER, outline=SIGNAL, width=4)
            draw.text((x + 30, 315), f"{i + 1:02}", font=pil_font(22, True), fill=CAUTION)
            draw.text((x + 30, 385), head, font=pil_font(34, True), fill=INK)
            draw.multiline_text((x + 30, 475), body, font=pil_font(25), fill=MUTED, spacing=12)
        path = FIG / filename
        image.save(path, quality=94)
        paths.append(path)
    return paths


def build_qmd(figures: list[Path]) -> Path:
    path = OUT / "lightguard_submission_report.qmd"
    lines = [
        "---\n",
        'title: "LightGuard"\n',
        'subtitle: "AMI 전력 흐름으로 가로등 점검 우선순위를 지원하는 운영 서비스"\n',
        'author: "김종백"\n',
        'date: "2026-08-21"\n',
        "lang: ko\nformat:\n  html:\n    css: report.css\n    toc: true\n---\n\n",
        f"> **핵심 정의**  \\n> {DEFINITION}\n\n",
        "## 한 페이지 요약\n\n",
    ]
    lines.extend(f"- **{value}** — {label} ({caveat})\n" for value, label, caveat in METRICS)
    lines.append(f"\n![서비스 흐름](figures/{figures[0].name})\n\n")
    for index, (title, deck, body, claim) in enumerate(SECTIONS, 1):
        lines.extend([f"## {index:02}. {title}\n\n", f"**{deck}**\n\n"])
        lines.extend(text + "\n\n" for text in body)
        lines.append(f"> **검증된 주장**: {claim}\n\n")
        if index == 5:
            lines.append(f"![증거 층](figures/{figures[1].name})\n\n")
        if index == 11:
            lines.append(f"![주장 경계](figures/{figures[2].name})\n\n")
    lines.append("## Appendix A. 금지 주장\n\n")
    lines.extend(f"- {item}\n" for item in DISCLAIMERS)
    lines.append("\n## Appendix B. 자료와 provenance\n\n")
    lines.extend(f"- {item}\n" for item in SOURCES)
    lines.append("\n## 제출 전 입력\n\n- 공식 신청서 팀명·대표자·연락처·E-MAIL\n- 공모분야 체크와 개인정보 동의 서명\n- 포털 파일명·용량·페이지 제한\n")
    path.write_text("".join(lines), encoding="utf-8")
    return path


def css() -> str:
    return f"""
:root{{--ink:{INK};--paper:{PAPER};--canvas:{CANVAS};--signal:{SIGNAL};--caution:{CAUTION};--line:{LINE};--muted:{MUTED}}}
*{{box-sizing:border-box}} body{{margin:0;background:var(--canvas);color:var(--ink);font-family:'NanumBarunGothic','Apple SD Gothic Neo',sans-serif;line-height:1.72}}
.cover{{min-height:78vh;display:grid;align-content:end;padding:clamp(36px,8vw,96px);color:white;background:linear-gradient(135deg,#102A43,#163A4D 62%,#0F766E);border-radius:0 0 44px 44px}}
.eyebrow{{color:#F7C948;font-size:.78rem;font-weight:800;letter-spacing:.14em}} h1{{font:800 clamp(3.2rem,10vw,7rem)/.94 'NanumSquare',sans-serif;letter-spacing:-.055em;margin:.25em 0}} .cover p{{max-width:760px;font-size:1.25rem;color:#E7F2EF}}
main{{width:min(100% - 32px,980px);margin:auto;padding:64px 0 120px}} .definition{{margin:0;padding:32px;background:var(--paper);border-left:7px solid var(--signal);font-size:1.2rem;font-weight:800}}
section{{padding:68px 0 28px;border-bottom:1px solid var(--line)}} .number{{color:var(--signal);font-weight:800;letter-spacing:.12em}} h2{{max-width:820px;font:800 clamp(2rem,5vw,3.4rem)/1.12 'NanumSquare',sans-serif;letter-spacing:-.045em;margin:.3em 0}} .deck{{max-width:760px;font-size:1.2rem;color:var(--muted)}} .copy{{max-width:760px}}
.claim{{margin:30px 0;padding:22px 26px;border:1px solid var(--signal);border-radius:18px;background:#E7F2EF;font-weight:800}} .metrics{{display:grid;grid-template-columns:repeat(3,1fr);gap:1px;border:1px solid var(--line);border-radius:20px;overflow:hidden;background:var(--line)}} .metric{{padding:24px;background:var(--paper)}} .metric strong{{display:block;color:var(--signal);font-size:2rem}} .metric small{{display:block;color:var(--muted);margin-top:8px}}
figure{{margin:42px 0}} figure img{{width:100%;border-radius:22px;box-shadow:0 16px 48px rgba(16,42,67,.12)}} figcaption{{color:var(--muted);font-size:.85rem}}
@media(max-width:700px){{.metrics{{grid-template-columns:1fr 1fr}}main{{width:min(100% - 24px,980px)}}}} @media(max-width:430px){{.metrics{{grid-template-columns:1fr}}.definition{{padding:22px}}}}
@media print{{@page{{size:A4;margin:16mm}}body{{background:white;font-size:10pt}}.cover{{min-height:250mm;break-after:page;border-radius:0}}main{{width:auto;padding:0}}section{{break-before:page;padding-top:8mm}}figure,.claim,.metrics{{break-inside:avoid}}}}
""".strip() + "\n"


def build_html(figures: list[Path]) -> Path:
    path = OUT / "LightGuard_Competition_Report.html"
    metric_html = "".join(f'<div class="metric"><strong>{html.escape(v)}</strong>{html.escape(l)}<small>{html.escape(c)}</small></div>' for v, l, c in METRICS)
    chunks = []
    for index, (title, deck, body, claim) in enumerate(SECTIONS, 1):
        copy = "".join(f"<p>{html.escape(p)}</p>" for p in body)
        figure = ""
        if index in (2, 5, 11):
            fig = figures[{2: 0, 5: 1, 11: 2}[index]]
            figure = f'<figure><img src="figures/{fig.name}" alt="{html.escape(title)}"><figcaption>출처: LightGuard v0.21 registry 및 공개자료. 해석 경계는 본문 참조.</figcaption></figure>'
        chunks.append(f'<section><div class="number">{index:02} / 12</div><h2>{html.escape(title)}</h2><p class="deck">{html.escape(deck)}</p><div class="copy">{copy}</div><div class="claim">검증된 주장 · {html.escape(claim)}</div>{figure}</section>')
    appendix = "".join(f"<li>{html.escape(x)}</li>" for x in DISCLAIMERS)
    sources = "".join(f"<li>{html.escape(x)}</li>" for x in SOURCES)
    page = f'''<!doctype html><html lang="ko"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>LightGuard Submission</title><link rel="stylesheet" href="report.css"></head><body><header class="cover"><div><div class="eyebrow">COMPETITION SUBMISSION · PUBLIC SERVICE</div><h1>LightGuard</h1><p>AMI 전력 흐름으로 가로등 점검 우선순위를 지원하는 운영 서비스</p><p><small>Submission Release v1.0 · 2026.08.21</small></p></div></header><main><blockquote class="definition">{html.escape(DEFINITION)}</blockquote><section><div class="number">EXECUTIVE SUMMARY</div><h2>이미 있는 AMI를 현장 확인의 두 번째 근거로</h2><p class="deck">센서를 더 설치하기 전에 자산·천문·운영자료와 AMI를 결합해 오늘 확인할 대상을 좁힙니다.</p><div class="metrics">{metric_html}</div></section>{''.join(chunks)}<section><div class="number">APPENDIX</div><h2>주장 경계와 provenance</h2><h3>금지 주장</h3><ul>{appendix}</ul><h3>자료</h3><ul>{sources}</ul><h3>제출 전 입력</h3><ul><li>팀명·대표자·연락처·E-MAIL</li><li>공모분야 체크와 개인정보 동의 서명</li><li>포털 파일명·용량·페이지 제한</li></ul></section></main></body></html>'''
    path.write_text(page, encoding="utf-8")
    return path


def prepare_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "NanumBarunGothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔바른고딕")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(16, 42, 67)
    normal.paragraph_format.line_spacing = 1.28
    for name, size, color in [("Title", 34, INK), ("Heading 1", 21, INK), ("Heading 2", 14, SIGNAL)]:
        style = doc.styles[name]
        style.font.name = "NanumSquareExtraBold"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔스퀘어 ExtraBold")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color[1:])


def build_docx(figures: list[Path]) -> Path:
    path = OUT / "LightGuard_Competition_Report.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(2)
    prepare_doc_styles(doc)
    header = section.header.paragraphs[0]
    header.text = "LIGHTGUARD  /  SUBMISSION RELEASE v1.0"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(82, 96, 109)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.add_run("COMPETITION SUBMISSION · PUBLIC SERVICE").font.color.rgb = RGBColor(15, 118, 110)
    doc.add_paragraph("LightGuard", style="Title")
    p = doc.add_paragraph("AMI 전력 흐름으로 가로등 점검 우선순위를 지원하는 운영 서비스")
    p.runs[0].font.size, p.runs[0].font.bold = Pt(18), True
    doc.add_paragraph("제안자 김종백  |  개인 참가  |  공익 서비스")
    doc.add_paragraph("Submission Release v1.0  |  2026.08.21")
    doc.add_picture(str(figures[0]), width=Inches(6.45))
    doc.add_page_break()
    doc.add_heading("한 페이지 요약", level=1)
    p = doc.add_paragraph(DEFINITION)
    p.runs[0].font.bold = True
    table = doc.add_table(rows=2, cols=3)
    table.style = "Light Shading Accent 1"
    for index, (value, label, caveat) in enumerate(METRICS):
        cell = table.cell(index // 3, index % 3)
        cell.text = ""
        p = cell.paragraphs[0]
        p.add_run(value + "\n").bold = True
        p.add_run(label + "\n")
        small = p.add_run(caveat)
        small.font.size, small.font.color.rgb = Pt(8), RGBColor(82, 96, 109)
    doc.add_heading("평가항목 대응", level=2)
    for item in [
        "문제 적합성: 실제 지자체의 대규모 확인 workload를 우선순위 문제로 정의",
        "차별성: 원격제어를 대체하지 않고 AMI 전력 흐름을 독립 확인",
        "실현 가능성: 수영구 204개 자산 객체와 Flutter 서비스 데모",
        "검증 신뢰성: 독립 QA, 32개 artifact 해시, registry와 금지주장 검사",
        "확장성: SIGNAL과 OPERATIONS layer를 지역별 데이터 수준에 맞게 결합",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()
    figure_map = {5: figures[1], 11: figures[2]}
    for index, (title, deck, body, claim) in enumerate(SECTIONS, 1):
        doc.add_heading(f"{index:02}. {title}", level=1)
        p = doc.add_paragraph(deck)
        p.runs[0].font.bold, p.runs[0].font.color.rgb = True, RGBColor(82, 96, 109)
        for text in body:
            doc.add_paragraph(text)
        p = doc.add_paragraph("검증된 주장  |  " + claim)
        p.runs[0].font.bold, p.runs[0].font.color.rgb = True, RGBColor(15, 118, 110)
        if index in figure_map:
            doc.add_picture(str(figure_map[index]), width=Inches(6.35))
        if index in (3, 6, 9):
            doc.add_page_break()
    doc.add_page_break()
    doc.add_heading("Appendix A. 금지 주장", level=1)
    for item in DISCLAIMERS:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Appendix B. 자료와 provenance", level=1)
    for item in SOURCES:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("제출 전 입력", level=2)
    doc.add_paragraph("공식 신청서의 팀명·대표자·연락처·E-MAIL, 공모분야 체크, 개인정보 동의 서명은 제출자가 직접 입력해야 합니다.")
    doc.save(path)
    return path


def build_pdf(figures: list[Path]) -> Path:
    path = OUT / "LightGuard_Competition_Report.pdf"
    pdfmetrics.registerFont(TTFont("Nanum", str(FONT)))
    pdfmetrics.registerFont(TTFont("NanumBold", str(FONT_BOLD)))
    pdfmetrics.registerFont(TTFont("NanumDisplay", str(FONT_DISPLAY)))
    sample = getSampleStyleSheet()
    body = ParagraphStyle("BodyKo", parent=sample["BodyText"], fontName="Nanum", fontSize=9.5, leading=15, textColor=colors.HexColor(INK), spaceAfter=7)
    h1 = ParagraphStyle("H1Ko", parent=sample["Heading1"], fontName="NanumDisplay", fontSize=19, leading=24, textColor=colors.HexColor(INK), spaceBefore=8, spaceAfter=9)
    deck = ParagraphStyle("DeckKo", parent=body, fontName="NanumBold", fontSize=11, leading=16, textColor=colors.HexColor(MUTED), spaceAfter=11)
    claim = ParagraphStyle("ClaimKo", parent=body, fontName="NanumBold", textColor=colors.HexColor(SIGNAL), backColor=colors.HexColor("#E7F2EF"), borderColor=colors.HexColor(SIGNAL), borderWidth=1, borderPadding=9, spaceBefore=7, spaceAfter=11)
    small = ParagraphStyle("SmallKo", parent=body, fontSize=7.5, leading=10, textColor=colors.HexColor(MUTED))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(LINE))
        canvas.line(18 * mm, 15 * mm, 192 * mm, 15 * mm)
        canvas.setFont("Nanum", 7)
        canvas.setFillColor(colors.HexColor(MUTED))
        canvas.drawString(18 * mm, 10 * mm, "LIGHTGUARD · SUBMISSION RELEASE v1.0")
        canvas.drawRightString(192 * mm, 10 * mm, str(doc.page))
        canvas.restoreState()

    story = [Spacer(1, 28 * mm), Paragraph("LightGuard", ParagraphStyle("TitleKo", fontName="NanumDisplay", fontSize=42, leading=48, textColor=colors.HexColor(INK))), Paragraph("AMI 전력 흐름으로 가로등 점검 우선순위를 지원하는 운영 서비스", ParagraphStyle("SubKo", parent=deck, fontSize=16, leading=23, textColor=colors.HexColor(SIGNAL))), Spacer(1, 12 * mm), Paragraph(DEFINITION, claim), Spacer(1, 10 * mm), PdfImage(str(figures[0]), width=174 * mm, height=98 * mm), PageBreak(), Paragraph("한 페이지 요약", h1)]
    metric_cells = [Paragraph(f"<font color='{SIGNAL}'><b>{v}</b></font><br/>{l}<br/><font size='7' color='{MUTED}'>{c}</font>", body) for v, l, c in METRICS]
    table = Table([metric_cells[:3], metric_cells[3:]], colWidths=[58 * mm] * 3, rowHeights=[33 * mm] * 2)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .5, colors.HexColor(LINE)), ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(PAPER)), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 8)]))
    story.extend([Paragraph(DEFINITION, claim), table, PageBreak()])
    figure_map = {5: figures[1], 11: figures[2]}
    for index, (title, deck_text, paragraphs, claim_text) in enumerate(SECTIONS, 1):
        story.extend([Paragraph(f"{index:02}. {title}", h1), Paragraph(deck_text, deck)])
        story.extend(Paragraph(text, body) for text in paragraphs)
        story.append(Paragraph("검증된 주장 · " + claim_text, claim))
        if index in figure_map:
            story.extend([PdfImage(str(figure_map[index]), width=174 * mm, height=98 * mm), Paragraph("출처: v0.21 registry 및 공개자료. 해석 경계는 본문 참조.", small)])
        story.append(Spacer(1, 5 * mm))
        if index in (3, 6, 9):
            story.append(PageBreak())
    story.extend([PageBreak(), Paragraph("Appendix A. 금지 주장", h1)])
    story.extend(Paragraph("• " + item, body) for item in DISCLAIMERS)
    story.append(Paragraph("Appendix B. 자료와 provenance", h1))
    story.extend(Paragraph("• " + item, body) for item in SOURCES)
    story.append(Paragraph("제출 전 입력", h1))
    story.append(Paragraph("공식 신청서의 팀명·대표자·연락처·E-MAIL, 공모분야 체크, 개인정보 동의 서명과 포털 제한을 확인해야 합니다.", body))
    pdf = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=20 * mm, title="LightGuard Competition Report")
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)
    return path


def build_visual_docx(pdf_path: Path) -> Path:
    """Create a visually stable DOCX from the reviewed PDF pages.

    Editable source remains available as QMD, HTML, and OFFICIAL_FORM_CONTENT.md.
    This avoids host-specific Korean font substitution in office renderers.
    """
    path = OUT / "LightGuard_Competition_Report.docx"
    with tempfile.TemporaryDirectory(prefix="lightguard-docx-pages-") as tmp:
        prefix = Path(tmp) / "page"
        subprocess.run(
            ["pdftoppm", "-png", "-r", "144", str(pdf_path), str(prefix)],
            check=True,
        )
        pages = sorted(Path(tmp).glob("page-*.png"), key=lambda item: int(item.stem.split("-")[-1]))
        doc = Document()
        doc.core_properties.title = "LightGuard Competition Report"
        doc.core_properties.author = "김종백"
        section = doc.sections[0]
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(1.5)
        section.left_margin = section.right_margin = Cm(1.5)
        for index, page in enumerate(pages):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(page), width=Cm(18))
            if index < len(pages) - 1:
                doc.add_page_break()
        doc.save(path)
    return path


def write_support_files() -> list[Path]:
    design = OUT / "DESIGN_CONTRACT.md"
    design.write_text(
        "# Design contract\n\n문제 → 왜 AMI → LightGuard → 실제 근거 → 운영 활용 → 한계 → 확장.\n\n"
        "- One claim per figure; insight-first title.\n- 수치·단위·지역·기간·출처·한계를 함께 표시.\n"
        "- SIGNAL·PLAUSIBILITY·OPERATIONS·PRODUCT 근거를 혼용하지 않음.\n"
        "- 앱 흐름은 이상징후 목록 → 이유 → 운영이력 → 원격관찰/현장점검.\n"
        "- 참고: OpenAI Editorial Infographic System, Reports/PDF Automation, quarto.report, Impeccable.\n",
        encoding="utf-8",
    )
    checklist = OUT / "RELEASE_CHECKLIST.md"
    checklist.write_text(
        "# Submission Release v1.0 Checklist\n\n"
        "- [x] submission/ Git 비추적\n- [x] v0.21 evidence artifact 비수정\n"
        "- [x] 주장·수치 caveat 표시\n- [x] 금지주장 명시\n- [x] QMD·HTML·PDF·DOCX·도표 생성\n"
        "- [x] 개인 참가자 김종백 표기\n- [ ] 연락처·E-MAIL 입력\n- [ ] 개인정보 동의 서명\n- [ ] 포털 용량·페이지 제한 확인\n",
        encoding="utf-8",
    )
    return [design, checklist]


def write_official_form_content(figures: list[Path]) -> Path:
    path = OUT / "OFFICIAL_FORM_CONTENT.md"
    path.write_text(
        f"""# LightGuard 공식 접수서류 입력 원고

## 신청자 정보

- 참가 형태: 개인
- 성명 / 대표자: 김종백
- 팀명(필수 입력란인 경우): 김종백
- 구성원: 해당 없음
- 연락처: [김종백 입력]
- E-MAIL: [김종백 입력]

## 제 목

**LightGuard: AMI 전력 흐름을 활용한 가로등 원격확인·현장점검 우선순위 지원 서비스**

## 공모분야

**공익 서비스**

## 추진배경

가로등 이상은 시민의 야간 안전과 이동 편의에 직접 연결되지만, 모든 분전함과 등기구를 같은 빈도로 현장 점검하기는 어렵다. 기존 원격제어 시스템은 점·소등 명령과 통신 상태를 관리하는 데 강점이 있으나, 명령 이후 실제 전력 흐름이 기대 범위에 들어왔는지를 독립적으로 확인하기는 어렵다. 대구공공시설관리공단의 고장등 관리자료 101,843건처럼 실제 지자체는 일상점검·직원신고·민원신고로 대규모 확인 업무를 수행한다. 따라서 필요한 것은 고장을 자동 확정한다는 과도한 모델이 아니라, 기존 AMI를 활용해 먼저 확인할 대상을 좁히는 설명 가능한 운영지원 서비스다.

**권장 이미지:** figures/{figures[2].name}  
**캡션:** LightGuard가 확인한 것과 확인하지 않은 것을 분리해 자동 고장확정, 확정 비용절감과 같은 과장 주장을 방지한다.

## 서비스 아이디어

LightGuard는 분전함별 자산정보, 예상 점등시간, 예상 정격부하, AMI 실측을 결합해 이상근거와 점검우선순위를 생성한다. 데이터 객체는 분전함 → 자산정보 → 예상 점등시간 → 예상 정격부하 → AMI 실측 → 이상근거 → 점검우선순위 순서로 고정된다. 예상보다 낮은 부하, 일출 이후에도 지속되는 소비 같은 신호를 탐지하되 최종 고장 여부는 담당자의 원격확인 또는 현장점검으로 확정한다. 데이터 결측은 데이터 품질 검토, 약한 신호는 원격관찰, 설명 가능하고 지속되는 신호는 현장점검 후보로 분기한다.

**권장 이미지:** figures/{figures[0].name}  
**캡션:** 예상과 AMI 관측의 차이를 이상근거로 설명하고 담당자의 다음 확인 행동으로 연결하는 서비스 흐름.

## 구체적 방법론

SIGNAL layer는 2026년 일출·일몰 및 시민박명, 분전함 정격부하, AMI 관측치의 시간·크기·지속시간·품질을 평가한다. OPERATIONS layer는 접수 당시 관측 가능한 접수경로, 계절, 자산별 과거 30/90/365일 사건, 이전 장기처리 이력, 열린 사건과 start-of-day backlog를 사용한다. 같은 날 사건 순서를 임의로 만들거나 미래 처리일을 가져오는 leakage는 금지한다. 실제 수영구 AMI가 없는 구간은 정격 3.4 kW 자산의 20% 부분소등, 일출 후 90분 점등 지속과 같은 scenario injection으로 판정 경로를 재현하며 실제 현장 정확도와 분리해 표시한다.

**권장 이미지:** figures/{figures[1].name}  
**캡션:** SIGNAL·PLAUSIBILITY·OPERATIONS·PRODUCT 근거는 서로 다른 질문에 답하며 하나의 정확도 숫자로 합치지 않는다.

## 결과물 및 성능

수영구 204개 분전함, 4,239등, 추정 정격부하 488.44 kW를 앱이 사용하는 객체로 구축했다. 통제된 이상 시나리오 46개를 탐지했고, 독립 QA PASS, 32개 release artifact 해시 일치, metric·claim·rubric registry 및 금지주장 검사를 통과했다. 대구 101,843건, 부여 3,437건, 울산 1,060건의 독립 운영자료를 통해 실제 유지관리 workload와 반복·처리 흐름을 확인했다. 단, 46개 시나리오 탐지는 실제 고장 정확도나 재현율이 아니며 운영자료도 AMI 성능의 대체 검증이 아니다.

**권장 이미지:** figures/{figures[3].name}  
**캡션:** 제출 전 검증 gate 결과. 46/46은 통제 시나리오 탐지이며 실제 현장 정확도가 아니다.

## 활용 방안

가로등 운영 담당자는 앱의 이상징후 목록에서 우선순위를 확인하고, 예상 부하·실측 신호·지속시간·판정근거를 검토한 뒤 반복 운영이력을 참고해 원격관찰 또는 현장점검을 선택한다. 도입 초기에는 기존 원격제어와 병행하는 읽기 전용 dashboard와 4~8주 shadow mode로 운영한다. 이후 담당자의 확인 결과를 feedback label로 축적해 지역별 threshold를 보정한다. 기존 제어 시스템을 교체하지 않으며, 분전함 AMI가 이미 있는 환경에서 낮은 추가설비 부담으로 적용한다.

**권장 이미지:** figures/{figures[4].name}  
**캡션:** 담당자 사용 예시. 목록에서 이유와 이력을 확인한 뒤 원격관찰 또는 현장점검으로 이어진다.

## 기대효과

LightGuard의 유형효과는 민원이 일정 비율 감소한다고 단정하는 것이 아니라, 실제 대규모 유지관리 workload에서 확인 대상을 좁혀 제한된 점검 시간을 설명 가능한 후보에 집중시키는 것이다. pilot에서는 후보 대비 원격 해소율, 현장점검 적중률, 사건당 확인시간, 반복 방문률, 장기 backlog와 데이터 품질 보류율을 사전 정의해 효과를 측정한다. 비용절감액과 민원 감소율은 실제 사용 전후 운영 로그가 확보된 뒤에만 산정한다.

## 전망 및 추후 과제

1. 실제 분전함 AMI와 자산 master를 읽기 전용으로 연결해 4~8주 shadow monitoring을 수행한다.
2. 담당자의 원격확인·현장점검 결과를 Gold/Silver label로 구조화한다.
3. 지역·계절을 넓혀 SIGNAL 안정성과 OPERATIONS local calibration을 분리 평가한다.
4. 개인정보와 원시 계측은 최소수집하고 release artifact는 해시와 registry로 관리한다.

## 이미지 사용 순서

1. 추진배경: 주장 경계
2. 서비스 아이디어: 서비스 플로우차트
3. 방법론: 증거 layer
4. 결과물 및 성능: 검증 성능 그래프
5. 활용 방안: 담당자 사용 예시

## 제출 시 주의

- 공식 양식의 셀 높이가 부족하면 본문은 위 원고를 60~70%로 축약하고 상세보고서를 별첨한다.
- 연락처·E-MAIL과 개인정보 동의 서명은 김종백 본인이 입력한다.
- 46/46 옆에는 반드시 통제 시나리오, 실제 현장 정확도 아님을 함께 표기한다.
- 비용절감액, 민원감소율, 자동 고장확정, 전국 무보정 적용 주장은 사용하지 않는다.
""",
        encoding="utf-8",
    )
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    figures = build_figures()
    qmd = build_qmd(figures)
    css_path = OUT / "report.css"
    css_path.write_text(css(), encoding="utf-8")
    html_path = build_html(figures)
    pdf_path = build_pdf(figures)
    docx_path = build_visual_docx(pdf_path)
    outputs = figures + [qmd, css_path, html_path, docx_path, pdf_path] + write_support_files()
    outputs.append(write_official_form_content(figures))
    sums = OUT / "SHA256SUMS.txt"
    sums.write_text("".join(f"{hashlib.sha256(path.read_bytes()).hexdigest()}  {path.relative_to(OUT)}\n" for path in sorted(outputs)), encoding="utf-8")
    print(f"Built {len(outputs)} artifacts in {OUT}")


if __name__ == "__main__":
    main()
