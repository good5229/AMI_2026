#!/usr/bin/env python3
"""Create a populated LightGuard application from the retained official DOCX."""

from __future__ import annotations

import shutil
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "release-v1.0" / "LightGuard_공모전_접수서류_김종백.docx"

TITLE = "LightGuard: AMI 전력 흐름 기반 가로등 점검 우선순위 지원 서비스"

FIELDS = {
    "추진배경": [
        "가로등 운영 과제: 대규모 확인 업무와 제한된 현장점검 자원",
        "기존 원격제어: 점·소등 명령 및 통신 상태 관리",
        "검증 공백: 명령 이후 실제 전력 흐름의 독립 확인 부족",
        "운영 근거: 대구 101,843건·부여 3,437건·울산 1,060건",
        "목표: 고장 자동 확정이 아닌 원격확인·현장점검 후보 축소",
    ],
    "서비스 아이디어": [
        "핵심 역할: 기존 AMI를 가로등 제어의 두 번째 확인자로 활용",
        "데이터 객체: 분전함 → 자산 → 예상 점등시간 → 예상 부하 → AMI → 이상근거 → 우선순위",
        "수영구 기반: 분전함 204개·가로등 4,239등·추정 정격부하 488.44 kW",
        "SIGNAL layer: 시민박명·정격부하·AMI·지속시간·품질 평가",
        "OPERATIONS layer: 반복 사건·열린 사건·start-of-day backlog 결합",
        "운영 분기: 데이터 품질 검토 / 원격관찰 / 현장점검 후보",
    ],
    "활용 방안": [
        "사용자: 지자체 도로조명·시설관리·관제·유지보수 담당자",
        "사용 흐름: 이상징후 목록 → 판정근거 → 운영이력 → 확인 행동",
        "도입 1단계: 기존 원격제어와 병행하는 읽기 전용 dashboard",
        "도입 2단계: 실제 AMI 기반 4~8주 shadow monitoring",
        "도입 3단계: 담당자 확인 결과 기반 지역별 threshold 보정",
    ],
    "기대효과": [
        "운영효과: 대규모 workload 내 우선 확인 후보 축소",
        "현장지원: 예상 부하·실측 신호·지속시간·반복이력의 통합 확인",
        "설비전략: 기존 AMI 우선 재사용·추가 센서의 선택적 도입",
        "pilot KPI: 원격 해소율·현장점검 적중률·확인시간·반복 방문률",
        "주장 경계: 자동 고장확정·민원감소율·확정 비용절감액 제외",
    ],
}


def norm(value: str) -> str:
    return " ".join(unicodedata.normalize("NFC", value).replace("\n", " ").split())


def unique_cells(row):
    seen = set()
    cells = []
    for cell in row.cells:
        marker = id(cell._tc)
        if marker not in seen:
            seen.add(marker)
            cells.append(cell)
    return cells


def style_run(run, size: float = 9.0, bold: bool = False, color: str = "102A43") -> None:
    run.font.name = "NanumBarunGothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔바른고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def fill_cell(cell, lines: list[str], size: float = 8.7, bullets: bool = True) -> None:
    cell.text = ""
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(("• " if bullets else "") + line)
        style_run(run, size=size, bold=not bullets)


def adjacent_cell(row, label_index: int):
    cells = unique_cells(row)
    if label_index + 1 < len(cells):
        return cells[label_index + 1]
    return None


def populate_row(row) -> None:
    cells = unique_cells(row)
    texts = [norm(cell.text) for cell in cells]
    for index, text in enumerate(texts):
        target = adjacent_cell(row, index)
        if target is None:
            continue
        if "제 목" in text or text == "제목":
            fill_cell(target, [TITLE], size=10.0, bullets=False)
        elif "공모분야" in text:
            fill_cell(
                target,
                ["☐ 국민편익    ☐ 신사업 및 Biz 모델    ☒ 공익 서비스"],
                size=9.3,
                bullets=False,
            )
        elif "추진배경" in text:
            fill_cell(target, FIELDS["추진배경"])
        elif "서비스 아이디어" in text:
            fill_cell(target, FIELDS["서비스 아이디어"], size=8.4)
        elif "활용 방안" in text or "활용방안" in text:
            fill_cell(target, FIELDS["활용 방안"])
        elif "기대효과" in text:
            fill_cell(target, FIELDS["기대효과"])
        elif "팀 명" in text or "팀명" in text or "단체명" in text:
            fill_cell(target, ["김종백 (개인 참가)"], size=9.5, bullets=False)
        elif "대표자" in text or text == "성명":
            fill_cell(target, ["김종백"], size=9.5, bullets=False)
        elif "구성원" in text:
            fill_cell(target, ["해당 없음"], size=9.0, bullets=False)
        elif "연락처" in text or "전화" in text:
            fill_cell(target, ["[김종백 직접 입력]"], size=9.0, bullets=False)
        elif "E-MAIL" in text.upper() or "이메일" in text:
            fill_cell(target, ["[김종백 직접 입력]"], size=9.0, bullets=False)


def replace_placeholders(doc: Document) -> None:
    replacements = {
        "010-0000-0000": "[연락처 직접 입력]",
        "000 000@0000.com": "[E-MAIL 직접 입력]",
        "000000@0000.com": "[E-MAIL 직접 입력]",
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in unique_cells(row):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for source, target in replacements.items():
                            if source in run.text:
                                run.text = run.text.replace(source, target)
                                style_run(run, size=8.5)


def populate_identity_and_consent(doc: Document) -> None:
    application = doc.tables[1]
    row0 = unique_cells(application.rows[0])
    fill_cell(row0[1], ["김종백 (개인 참가)"], size=9.5, bullets=False)

    leader = unique_cells(application.rows[2])
    fill_cell(leader[1], ["개인"], size=9.0, bullets=False)
    fill_cell(leader[2], ["김종백"], size=9.5, bullets=False)
    fill_cell(leader[3], ["[직접 입력]"], size=8.5, bullets=False)
    fill_cell(leader[4], ["[직접 입력]"], size=8.5, bullets=False)

    for row_index in (3, 4):
        member = unique_cells(application.rows[row_index])
        fill_cell(member[1], ["해당 없음"], size=8.5, bullets=False)
        fill_cell(member[2], ["해당 없음"], size=8.5, bullets=False)
        fill_cell(member[3], ["-"], size=8.5, bullets=False)
        fill_cell(member[4], ["-"], size=8.5, bullets=False)

    agreement = application.cell(5, 0)
    fill_cell(
        agreement,
        [
            "상기 신청자는 아래 유의 사항에 동의합니다.",
            "2026. [월] [일]    신청자 김종백    (인/서명)",
        ],
        size=9.0,
        bullets=False,
    )

    consent_cell = doc.tables[2].cell(2, 0)
    for paragraph in consent_cell.paragraphs:
        if "위 내용에 동의합니다" in norm(paragraph.text):
            paragraph.text = "( 위 내용에 동의합니다. ☒ 동의함 ☐ 동의하지 않음 )"
            for run in paragraph.runs:
                style_run(run, size=9.0, bold=True)

    signature_rows = [
        paragraph
        for paragraph in doc.paragraphs
        if "소 속" in norm(paragraph.text) and "성 명" in norm(paragraph.text)
    ]
    if signature_rows:
        signature_rows[0].text = "소 속 : 개인 참가    성 명 : 김종백    (인/서명)"
        for run in signature_rows[0].runs:
            style_run(run, size=9.5)
        for paragraph in signature_rows[1:]:
            paragraph.text = ""
    for paragraph in doc.paragraphs:
        if norm(paragraph.text) == "2026. . .":
            paragraph.text = "2026. [월] [일]"
            for run in paragraph.runs:
                style_run(run, size=10.0, bold=True)


def main() -> None:
    templates = sorted((ROOT / "official_docs").glob("*MS*docx"))
    if not templates:
        raise FileNotFoundError("Official MS Word application template not found")
    template = templates[0]
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, OUT)
    doc = Document(OUT)
    doc.core_properties.title = "LightGuard 공모전 접수서류"
    doc.core_properties.author = "김종백"
    for table in doc.tables:
        for row in table.rows:
            populate_row(row)
    replace_placeholders(doc)
    populate_identity_and_consent(doc)
    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        if "2026." in text and "(" in text and ")" in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
